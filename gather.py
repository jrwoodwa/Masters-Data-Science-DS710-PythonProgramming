"""
ds 710: Programming for Data Science
Project: Job Search Research

John Woodward

Do not rename this file
"""

# ----------------------------------------

## libraries import
# import os
# import matplotlib.pyplot as plt
# import seaborn as sns
# import numpy as np
import requests # web crawling
import pandas as pd # data manipulation
from datetime import datetime, date
import docx # for creating a word document

# prior code import

def filename_checker(ext,name):
    '''
    returns the filename with an extension

    Parameters
    ----------
    ext : str
        extension of the filename.
    name : str
        descriptive filename.

    Returns
    -------
    return_filename : str
        filename with extension.

    '''
    return name if name.endswith(ext) else name + '.' + ext

# ----------------------------------------

# Task 1: Select and gather data

## Subtask 1.1: gather.py

url1 = 'https://en.wikipedia.org/wiki/List_of_U.S._states_by_adjusted_per_capita_personal_income'
url2 = 'https://en.wikipedia.org/wiki/List_of_states_and_territories_of_the_United_States'

def html_wikitable_to_dataframe(url):
    '''
    Take a wikipedia table and output the pandas dataframes

    Parameters
    ----------
    url : str
        Wikipedia string URL to see tables.

    Returns
    -------
    dfs : pandas data frames
        The tables are in an array, still need to index afterward which table.

    '''
    # Request the web page and get the response
    response = requests.get(url)
    
    # Get the HTML source code of the web page
    html = response.content.decode()
    
    # Use pandas to read the HTML table
    dfs = pd.read_html(html, 
                       attrs={'class': 'wikitable'})
    return dfs

purchasepower_df = html_wikitable_to_dataframe(url=url1)[0]

# Keep only the columns `State` and `Purchasing...`
purchasepower_df = purchasepower_df[['State', 'Purchasing power of $1.00 (2019)[5]']]

# Renaming the column
purchasepower_df = purchasepower_df.rename(columns={'Purchasing power of $1.00 (2019)[5]': 'Purchase_Power_2019','State':'State_Long'})

# use another wikipedia table to gather the us state abbreviations
abbreviations_df = html_wikitable_to_dataframe(url=url2)[0]

# select only the first two columns and rename them
abbreviations_df = abbreviations_df.iloc[:, :2]
abbreviations_df.columns = ['State_Long', 'State']

# Clean up the State_Long column by removing "[B]"
abbreviations_df['State_Long'] = abbreviations_df['State_Long'].str.replace('\[.*\]', '', 
                                                                            regex=True)

merged_df = purchasepower_df.merge(abbreviations_df, 
                                   on='State_Long', 
                                   how='inner')

columns_order = ['State_Long', 'State', 'Purchase_Power_2019']
merged_df = merged_df[columns_order]

#merged_df

# Save to a csv file
merged_df.to_csv('purchase_power_by_state.csv', 
                 index=False)

# ---------------------------------------

## Subtask 1.2: data_provenance.pdf

# Save a provenance WORD document (still need to open manually and save as PDF)

# Set up the document
doc = docx.Document()
doc.add_heading('Data Provenance - John Woodward', 0)

doc.add_heading('Data', level=1)

# URL 1
TOS1 = 'https://foundation.wikimedia.org/wiki/Policy:Terms_of_Use/en#4._Refraining_from_Certain_Activities'
doc.add_heading('U.S. states by purchase power', level=2)
doc.add_paragraph(f"Description: {'This Wikipedia page provides a ranked list of U.S. states by their adjusted per capita personal income, which considers the cost of living differences between states.'}")
doc.add_paragraph(f"Date of Download: {datetime.today().strftime('%Y-%m-%d')}")
doc.add_paragraph(f'Source: {url1}')
doc.add_paragraph(f'Terms of Service: {TOS1}')
doc.add_paragraph('License: Creative Commons Attribution-ShareAlike 3.0 Unported License (CC BY-SA 3.0)')

# URL 2
doc.add_heading('U.S. states and abbreviations', level=2)
doc.add_paragraph(f"Description: {'This Wikipedia page lists all 50 states, with brief descriptions and related information.'}")
doc.add_paragraph(f"Date of Download: {datetime.today().strftime('%Y-%m-%d')}")
doc.add_paragraph(f'Source: {url2}')
doc.add_paragraph(f'Terms of Service: {TOS1}')
doc.add_paragraph('License: Creative Commons Attribution-ShareAlike 3.0 Unported License (CC BY-SA 3.0)')

# LinkedIn
TOS2 = 'https://www.linkedin.com/legal/user-agreement#dos'
url3 = 'https://www.linkedin.com/jobs/search/?currentJobId=3518862183&distance=25&geoId=103644278&keywords=(cplex%20OR%20%E2%80%9Clinear%20programming%E2%80%9D%20OR%20gurobi)%20and%20(python%20sql)'
description3 = 'I used the URL provided to conduct job searches on LinkedIn, with the search criteria being: (cplex OR "linear programming" OR gurobi) and (python sql). \n\nFollowing this, a Lix browser plug-in was downloaded https://chrome.google.com/webstore/detail/lix-linkedin-scraping-ema/ceplokfhfeekddamgoaojabdhkggafnk, registered at https://lix-it.com/login, and linked with an active LinkedIn account. \n\nSubsequently, I applied the Lix-it tool to obtain high-level meta-data for the job listings (e.g., Organization, Job Title). \n\nLastly, I manually downloaded the lower-level job description data by hiring a Fiverr contractor for $50 to comply with the terms and conditions prohibiting automated web scraping. The process took four hours of manual work, was one-time, and performed non-commercially.\n\nIn retrospect, the Lix tool gathering metadata may have skirted terms of service for LinkedIn, but according to the additional context provided, this is considered acceptable.'
additional = 'https://www.zdnet.com/article/court-rules-that-data-scraping-is-legal-in-linkedin-appeal/'
additional2 = 'https://www.eff.org/deeplinks/2020/04/federal-judge-rules-it-not-crime-violate-websites-terms-service'
additional3 = 'https://www.eff.org/deeplinks/2022/04/scraping-public-websites-still-isnt-crime-court-appeals-declares'
doc.add_heading('Job Search results', level=2)
doc.add_paragraph(f"Description: {description3}")
doc.add_paragraph(f"Date of Download: {date(2023, 4, 11).strftime('%Y-%m-%d')}")
doc.add_paragraph(f'Source: {url3}')
doc.add_paragraph(f'Terms of Service: {TOS2}')
doc.add_paragraph('License: N/A')
doc.add_paragraph(f'Additional Context: {additional} \n {additional2} \n {additional3}')

# Save the document
doc.save('data_provenance.docx')

# ------------------------------

if __name__ == "__main__":
    pass